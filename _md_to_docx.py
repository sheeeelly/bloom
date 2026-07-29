# -*- coding: utf-8 -*-
"""Convert 智能选品 一期 方案.md to Word (simple markdown subset)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "智能选品 一期 方案.md"
OUT_PATH = ROOT / "智能选品 一期 方案_v1.2_终稿.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, style=None, size=11, bold=False, space_after=6):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    # strip highlight markers ==...==
    text = re.sub(r"==([^=]+)==", r"\1", text)
    # inline code `...`
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, size=size)
            run.font.name = "Consolas"
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def is_table_sep(line: str) -> bool:
    s = line.strip()
    return bool(s.startswith("|") and re.match(r"^\|[\s\-:|]+\|$", s))


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return [re.sub(r"\*\*|==|`", "", c) for c in cells]


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = row[j] if j < len(row) else ""
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=9, bold=(i == 0))
    doc.add_paragraph()


def convert():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    # title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智能选品一期方案（v1.2 访谈校准终稿）")
    set_run_font(run, size=18, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("访谈来源：2026-07-20 苗诗雨（ACC）｜含待群装/BD 补访项")
    set_run_font(r2, size=10, color=RGBColor(0x66, 0x66, 0x66))

    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        if raw.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_paragraph(doc, raw, size=9)
            i += 1
            continue

        # skip image lines (zoom signed URLs not useful in Word)
        if raw.startswith("![") and "](" in raw:
            add_paragraph(doc, "（此处原文档含截图/附件，请参阅 Markdown 版）", size=9)
            i += 1
            continue

        # table block
        if raw.strip().startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            rows = [parse_table_row(raw)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_sep(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if not raw.strip():
            i += 1
            continue

        if raw.startswith("# "):
            add_paragraph(doc, raw[2:], size=16, bold=True, space_after=10)
        elif raw.startswith("## "):
            add_paragraph(doc, raw[3:].replace("**", ""), size=14, bold=True, space_after=8)
        elif raw.startswith("### "):
            add_paragraph(doc, raw[4:].replace("**", ""), size=12, bold=True, space_after=6)
        elif raw.startswith("#### "):
            add_paragraph(doc, raw[5:].replace("**", ""), size=11, bold=True, space_after=4)
        elif raw.startswith("> "):
            p = add_paragraph(doc, raw[2:], size=10)
            p.paragraph_format.left_indent = Pt(12)
        elif raw.startswith("- ") or raw.startswith("· ") or re.match(r"^\d+\.\s", raw):
            add_paragraph(doc, raw, size=11)
        else:
            add_paragraph(doc, raw, size=11)
        i += 1

    doc.save(OUT_PATH)
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    convert()
